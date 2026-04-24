from __future__ import annotations

import argparse
import json
import os
import re
from copy import deepcopy
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.parts.numbering import NumberingPart
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


# ============================================================
# --- config & constants ---
# ============================================================

OUTPUT_SUFFIX = "_f"
DEFAULT_CONFIG_PATH = Path(__file__).with_name("config.json")

STYLE_BODY = "body"
STYLE_H1 = "h1"
STYLE_H2 = "h2"
STYLE_H3 = "h3"

H1_PREFIXES = "一二三四五六七八九十"

INTRO_TITLE_TEXT = "引言"
CONCLUSION_TITLE_TEXT = "结语"
REFERENCES_TITLE_TEXT = "参考文献"
ACK_TITLE_TEXT = "致谢"

REFERENCES_BODY_RULE = {
    "east_asia_font": "宋体",
    "size_pt": 10.5,
    "bold": False,
    "line_spacing": 1.5,
    "alignment": "left",
    "hanging_indent_chars": 2,
}


def load_config(config_path: Path) -> dict:
    if not config_path.exists():
        raise FileNotFoundError(f"Config file not found: {config_path}")
    return json.loads(config_path.read_text(encoding="utf-8"))


def _debug_enabled(config: dict | None) -> bool:
    if os.environ.get("PAPER_FORMAT_DEBUG", "").strip().lower() in {"1", "true", "yes", "on"}:
        return True
    if not config:
        return False
    debug_config = config.get("debug")
    if isinstance(debug_config, bool):
        return debug_config
    if isinstance(debug_config, dict):
        return bool(debug_config.get("enabled", False))
    return False


def _debug_log(config: dict | None, message: str) -> None:
    if _debug_enabled(config):
        print(f"[DEBUG] {message}")


# ============================================================
# --- low-level docx helpers ---
# ============================================================

def _resolve_alignment(alignment: str | None) -> WD_PARAGRAPH_ALIGNMENT | None:
    if not alignment:
        return None
    mapping = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "distributed": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE,
        "distribute": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE,
    }
    return mapping.get(alignment.lower())


def _set_run_font(run, east_asia_font: str, size_pt: float, bold: bool) -> None:
    run.font.name = "Times New Roman"
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def _clear_indent(pf) -> None:
    pf.first_line_indent = None
    pf.left_indent = None


def apply_paragraph_rule(paragraph, rule: dict) -> None:
    pf = paragraph.paragraph_format
    size_pt = float(rule.get("size_pt", 12))
    _clear_indent(pf)

    if "first_line_indent_chars" in rule:
        pf.first_line_indent = Pt(size_pt * int(rule["first_line_indent_chars"]))
    if "hanging_indent_chars" in rule:
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(size_pt * int(rule["hanging_indent_chars"]))

    if "line_spacing_pt" in rule:
        pf.line_spacing = Pt(float(rule["line_spacing_pt"]))
    elif "line_spacing" in rule:
        pf.line_spacing = float(rule["line_spacing"])

    if "space_before_lines" in rule:
        pf.space_before = Pt(size_pt * float(rule["space_before_lines"]))
    if "space_after_lines" in rule:
        pf.space_after = Pt(size_pt * float(rule["space_after_lines"]))

    para_alignment = _resolve_alignment(rule.get("alignment"))
    if para_alignment is not None:
        paragraph.alignment = para_alignment

    for run in paragraph.runs:
        _set_run_font(
            run,
            str(rule.get("east_asia_font", "宋体")),
            size_pt,
            bool(rule.get("bold", False)),
        )


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _prev_paragraph(paragraph) -> Paragraph | None:
    element = paragraph._p.getprevious()
    while element is not None and element.tag != qn("w:p"):
        element = element.getprevious()
    if element is None:
        return None
    return Paragraph(element, paragraph._parent)


def _next_paragraph(paragraph) -> Paragraph | None:
    element = paragraph._p.getnext()
    while element is not None and element.tag != qn("w:p"):
        element = element.getnext()
    if element is None:
        return None
    return Paragraph(element, paragraph._parent)


def _is_blank_paragraph(paragraph) -> bool:
    # 仅含换页符的段落文本为空，但属于结构内容，不视作空行。
    if paragraph._p is None:
        return True
    if paragraph._p.xpath(".//w:br[@w:type='page']"):
        return False
    return not paragraph.text.strip()


def _insert_blank_paragraph_before(paragraph) -> Paragraph:
    p = OxmlElement("w:p")
    paragraph._p.addprevious(p)
    return Paragraph(p, paragraph._parent)


def _insert_blank_paragraph_after(paragraph) -> Paragraph:
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    return Paragraph(p, paragraph._parent)


def _delete_paragraph(paragraph) -> None:
    # 仅把节点从 XML 树上摘除；保留 Paragraph 包装对象的 _p 引用，
    # 这样调用方对已删段落再做只读访问（比如 .text）也不会炸。
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _paragraph_has_page_break(paragraph) -> bool:
    if paragraph._p is None:
        return False
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def _ensure_page_break_before(paragraph) -> None:
    prev = _prev_paragraph(paragraph)
    if prev is None:
        return
    if _paragraph_has_page_break(prev):
        return
    prev.add_run().add_break(WD_BREAK.PAGE)


def _make_blank_body_paragraph(blank: Paragraph, body_rule: dict) -> None:
    apply_paragraph_rule(blank, body_rule)
    blank.paragraph_format.first_line_indent = Pt(0)


def _ensure_blank_before(paragraph, body_rule: dict) -> None:
    prev = _prev_paragraph(paragraph)
    if prev is not None and _is_blank_paragraph(prev):
        return
    blank = _insert_blank_paragraph_before(paragraph)
    _make_blank_body_paragraph(blank, body_rule)


def _ensure_blank_after(paragraph, body_rule: dict) -> None:
    nxt = _next_paragraph(paragraph)
    if nxt is not None and _is_blank_paragraph(nxt):
        return
    blank = _insert_blank_paragraph_after(paragraph)
    _make_blank_body_paragraph(blank, body_rule)


def _remove_blanks_before(paragraph) -> None:
    while True:
        prev = _prev_paragraph(paragraph)
        if prev is None or not _is_blank_paragraph(prev):
            return
        _delete_paragraph(prev)


def _remove_blanks_after(paragraph) -> None:
    while True:
        nxt = _next_paragraph(paragraph)
        if nxt is None or not _is_blank_paragraph(nxt):
            return
        _delete_paragraph(nxt)


def _collapse_blanks_before(paragraph, max_count: int = 1) -> None:
    """把标题前的连续空行压缩到不超过 max_count 行。"""
    while True:
        prev = _prev_paragraph(paragraph)
        if prev is None or not _is_blank_paragraph(prev):
            return
        prev_prev = _prev_paragraph(prev)
        if prev_prev is None or not _is_blank_paragraph(prev_prev):
            return
        _delete_paragraph(prev_prev)


def _collapse_blanks_after(paragraph, max_count: int = 1) -> None:
    while True:
        nxt = _next_paragraph(paragraph)
        if nxt is None or not _is_blank_paragraph(nxt):
            return
        nxt_nxt = _next_paragraph(nxt)
        if nxt_nxt is None or not _is_blank_paragraph(nxt_nxt):
            return
        _delete_paragraph(nxt_nxt)


def _set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _add_page_number_run(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._element.append(fld_begin)
    run._element.append(instr_text)
    run._element.append(fld_end)


def _remove_footer_references(section) -> None:
    sect_pr = section._sectPr
    for node in list(sect_pr.findall(qn("w:footerReference"))):
        sect_pr.remove(node)


def _clear_footer(footer) -> None:
    for paragraph in footer.paragraphs:
        paragraph.text = ""


def _get_or_create_numbering_part(document: Document) -> NumberingPart:
    try:
        return document.part.numbering_part
    except NotImplementedError:
        package = document.part.package
        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            CT.WML_NUMBERING,
            parse_xml(
                '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            ),
            package,
        )
        document.part.relate_to(numbering_part, RT.NUMBERING)
        return numbering_part


def _create_heading_numbering(document: Document) -> int:
    numbering_part = _get_or_create_numbering_part(document)
    numbering = numbering_part.numbering_definitions._numbering
    abstract_nums = numbering.xpath("./w:abstractNum")
    nums = numbering.xpath("./w:num")
    next_abstract_id = (
        max((int(node.get(qn("w:abstractNumId"))) for node in abstract_nums), default=-1) + 1
    )
    next_num_id = max((int(node.get(qn("w:numId"))) for node in nums), default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(next_abstract_id))

    def _append_numbering_rpr(lvl, size_pt: float) -> None:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:cs"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), "Times New Roman")
        r_pr.append(r_fonts)

        size_half_points = str(int(round(size_pt * 2)))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), size_half_points)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), size_half_points)
        r_pr.extend([sz, sz_cs])
        lvl.append(r_pr)

    # 一级：一、二、三、
    lvl0 = OxmlElement("w:lvl")
    lvl0.set(qn("w:ilvl"), "0")
    start0 = OxmlElement("w:start")
    start0.set(qn("w:val"), "1")
    numfmt0 = OxmlElement("w:numFmt")
    numfmt0.set(qn("w:val"), "chineseCounting")
    lvltext0 = OxmlElement("w:lvlText")
    lvltext0.set(qn("w:val"), "%1、")
    suff0 = OxmlElement("w:suff")
    suff0.set(qn("w:val"), "space")
    lvl0.extend([start0, numfmt0, lvltext0, suff0])
    _append_numbering_rpr(lvl0, 15.0)

    # 二级：（一）（二）（三）
    lvl1 = OxmlElement("w:lvl")
    lvl1.set(qn("w:ilvl"), "1")
    start1 = OxmlElement("w:start")
    start1.set(qn("w:val"), "1")
    numfmt1 = OxmlElement("w:numFmt")
    numfmt1.set(qn("w:val"), "chineseCounting")
    lvltext1 = OxmlElement("w:lvlText")
    lvltext1.set(qn("w:val"), "（%2）")
    suff1 = OxmlElement("w:suff")
    suff1.set(qn("w:val"), "space")
    lvl1.extend([start1, numfmt1, lvltext1, suff1])
    _append_numbering_rpr(lvl1, 14.0)

    # 三级：1. 2. 3.
    lvl2 = OxmlElement("w:lvl")
    lvl2.set(qn("w:ilvl"), "2")
    start2 = OxmlElement("w:start")
    start2.set(qn("w:val"), "1")
    numfmt2 = OxmlElement("w:numFmt")
    numfmt2.set(qn("w:val"), "decimal")
    lvltext2 = OxmlElement("w:lvlText")
    lvltext2.set(qn("w:val"), "%3.")
    suff2 = OxmlElement("w:suff")
    suff2.set(qn("w:val"), "space")
    lvl2.extend([start2, numfmt2, lvltext2, suff2])
    _append_numbering_rpr(lvl2, 12.0)

    abstract_num.extend([lvl0, lvl1, lvl2])
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(next_abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return next_num_id


def _set_heading_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.numPr
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))

    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


# ============================================================
# --- landmark detection & segmentation ---
# ============================================================

class Landmark(Enum):
    INTRO = "intro"
    CONCLUSION = "conclusion"
    REFERENCES = "references"
    ACK = "ack"


def _normalize_title_token(text: str) -> str:
    token = text.strip()
    token = re.sub(r"^[（(]\s*[一二三四五六七八九十]+\s*[)）]", "", token)
    token = re.sub(r"^[一二三四五六七八九十]+、", "", token)
    token = re.sub(r"^\d+(?:\.\d+)*\.", "", token)
    token = re.sub(r"[\s·•\\-—_<>《》【】\\[\\]()（）:：,，。；;!！?？\"'`]+", "", token)
    return token


def _strip_number_prefix(text: str) -> str:
    candidate = text.strip()
    candidate = re.sub(r"^[一二三四五六七八九十]+、", "", candidate)
    candidate = re.sub(r"^（[一二三四五六七八九十]+）", "", candidate)
    candidate = re.sub(r"^\(\s*[一二三四五六七八九十]+\s*\)", "", candidate)
    candidate = re.sub(r"^\d+(?:\.\d+)*\.", "", candidate)
    return candidate.strip()


def _reference_title_tokens(config: dict) -> set[str]:
    structure_config = config.get("structure", {})
    configured = _normalize_title_token(
        structure_config.get("references_title", REFERENCES_TITLE_TEXT)
    )
    tokens = {REFERENCES_TITLE_TEXT, "引用", "引用页"}
    if configured:
        tokens.add(configured)
    return {_normalize_title_token(token) for token in tokens if token}


def _compact_stripped(text: str) -> str:
    return re.sub(r"\s+", "", _strip_number_prefix(text))


def _paragraph_debug_summary(paragraph) -> str:
    text = paragraph.text.replace("\n", "\\n")
    style_name = (paragraph.style.name or "").strip()
    has_field = bool(paragraph._p.xpath(".//w:fldSimple")) or bool(paragraph._p.xpath(".//w:instrText"))
    return f"text='{text}', style='{style_name}', has_field={has_field}"


def detect_landmark(paragraph, config: dict) -> Landmark | None:
    """唯一真值源：把一个段落识别为某个里程碑，或返回 None。"""
    text = paragraph.text
    if not text.strip():
        return None
    normalized = _normalize_title_token(text)
    compact = _compact_stripped(text)

    if compact == INTRO_TITLE_TEXT or normalized == INTRO_TITLE_TEXT:
        return Landmark.INTRO
    if compact == CONCLUSION_TITLE_TEXT or normalized == CONCLUSION_TITLE_TEXT:
        return Landmark.CONCLUSION
    if normalized in _reference_title_tokens(config):
        return Landmark.REFERENCES
    if normalized == ACK_TITLE_TEXT:
        return Landmark.ACK
    return None


@dataclass
class DocumentSegments:
    """把整份文档按里程碑切分后各段的段落引用。每段互不依赖。"""

    document: Document
    intro_title: Paragraph | None = None
    conclusion_title: Paragraph | None = None
    body_paragraphs: list[Paragraph] = field(default_factory=list)
    references_title: Paragraph | None = None
    references_body: list[Paragraph] = field(default_factory=list)
    ack_title: Paragraph | None = None
    ack_body: list[Paragraph] = field(default_factory=list)


@dataclass
class FormatReport:
    total_paragraphs: int = 0
    landmarks: list[str] = field(default_factory=list)
    body_titles: int = 0
    body_text: int = 0
    figure_captions: int = 0
    references_text: int = 0
    acknowledgment_text: int = 0
    section_titles: int = 0
    margins_updated: int = 0
    page_numbers_updated: int = 0

    def to_dict(self) -> dict:
        return {
            "total_paragraphs": self.total_paragraphs,
            "landmarks": self.landmarks,
            "body_titles": self.body_titles,
            "body_text": self.body_text,
            "figure_captions": self.figure_captions,
            "references_text": self.references_text,
            "acknowledgment_text": self.acknowledgment_text,
            "section_titles": self.section_titles,
            "margins_updated": self.margins_updated,
            "page_numbers_updated": self.page_numbers_updated,
        }


def scan_segments(document: Document, config: dict) -> DocumentSegments:
    paragraphs = list(document.paragraphs)
    landmarks: dict[Landmark, int] = {}
    for idx, paragraph in enumerate(paragraphs):
        lm = detect_landmark(paragraph, config)
        # 只记录第一次出现，避免正文里偶然出现 "引言" 字样干扰。
        if lm is not None and lm not in landmarks:
            landmarks[lm] = idx
            _debug_log(
                config,
                f"landmark[{lm.value}] at idx={idx}: {_paragraph_debug_summary(paragraph)}",
            )

    intro_idx = landmarks.get(Landmark.INTRO)
    conclusion_idx = landmarks.get(Landmark.CONCLUSION)
    refs_idx = landmarks.get(Landmark.REFERENCES)
    ack_idx = landmarks.get(Landmark.ACK)

    segments = DocumentSegments(document=document)

    # Body 段：从 引言 到 参考文献/致谢 之前（含结语）
    if intro_idx is not None:
        segments.intro_title = paragraphs[intro_idx]
        body_end = min(
            (i for i in (refs_idx, ack_idx) if i is not None and i > intro_idx),
            default=len(paragraphs),
        )
        segments.body_paragraphs = paragraphs[intro_idx:body_end]
        if conclusion_idx is not None and intro_idx < conclusion_idx < body_end:
            segments.conclusion_title = paragraphs[conclusion_idx]
    elif refs_idx is None and ack_idx is None:
        # 兼容遗留文档：完全没有里程碑时，把首个非空段落之后当作正文。
        first_non_empty = next(
            (i for i, p in enumerate(paragraphs) if p.text.strip()), None
        )
        if first_non_empty is not None:
            segments.body_paragraphs = paragraphs[first_non_empty:]

    # References 段
    if refs_idx is not None:
        segments.references_title = paragraphs[refs_idx]
        refs_end = ack_idx if ack_idx is not None and ack_idx > refs_idx else len(paragraphs)
        segments.references_body = paragraphs[refs_idx + 1 : refs_end]

    # Ack 段
    if ack_idx is not None:
        segments.ack_title = paragraphs[ack_idx]
        segments.ack_body = paragraphs[ack_idx + 1 :]

    return segments


# ============================================================
# --- segment processors ---
# ============================================================

def _heading_level_by_prefix(text: str) -> int | None:
    t = text.strip()
    if re.match(r"^[一二三四五六七八九十]+、", t):
        return 1
    if re.match(r"^(（[一二三四五六七八九十]+）|\(\s*[一二三四五六七八九十]+\s*\))", t):
        return 2
    if re.match(r"^\d+(?:\.\d+)*\.", t):
        return 3
    return None


def _resolve_style_rule_key(paragraph) -> str:
    style_name = (paragraph.style.name or "").lower()
    if "heading 1" in style_name or "标题 1" in style_name:
        return STYLE_H1
    if "heading 2" in style_name or "标题 2" in style_name:
        return STYLE_H2
    if "heading 3" in style_name or "标题 3" in style_name:
        return STYLE_H3
    return STYLE_BODY


def _looks_like_heading(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    style_key = _resolve_style_rule_key(paragraph)
    if style_key in {STYLE_H1, STYLE_H2, STYLE_H3}:
        return True
    if _heading_level_by_prefix(text) is not None:
        return True
    if len(text) <= 24 and not any(ch in text for ch in "，。；：！？,.?!:;"):
        return True
    return False


def _resolve_heading_level(paragraph) -> int:
    text = paragraph.text.strip()
    level = _heading_level_by_prefix(text)
    if level is not None:
        return level
    style_key = _resolve_style_rule_key(paragraph)
    if style_key == STYLE_H2:
        return 2
    if style_key == STYLE_H3:
        return 3
    return 1


def _is_figure_caption(text: str) -> bool:
    # 仅识别 "图N" 开头的段落，不强制 "图N-N" 形式。
    return bool(re.match(r"^图\s*\d+", text.strip()))


def _apply_figure_caption(paragraph, body_rule: dict) -> None:
    apply_paragraph_rule(paragraph, body_rule)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    for run in paragraph.runs:
        _set_run_font(run, "黑体", 10.5, False)


def _apply_page_title_style(paragraph, rule: dict) -> None:
    """应用独立成页的标题样式：套规则 + 强制 space_before=0（因为前面是换页）。"""
    apply_paragraph_rule(paragraph, rule)
    paragraph.paragraph_format.space_before = Pt(0)


# ---------- Body (引言 -> 结语) ----------

def process_body(segments: DocumentSegments, config: dict, report: FormatReport | None = None) -> None:
    """独立处理正文段：引言 / H1 / H2 / H3 / 图序 / 正文 / 结语，含段内空行规则。"""
    if not segments.body_paragraphs:
        return

    style_config = config["styles"]
    body_rule = style_config[STYLE_BODY]
    h1_rule = style_config[STYLE_H1]
    h2_rule = style_config[STYLE_H2]
    h3_rule = style_config[STYLE_H3]

    heading_num_id = _create_heading_numbering(segments.document)

    chapter_idx = 0
    section_idx = 0
    sub_idx = 0

    intro_title = segments.intro_title
    conclusion_title = segments.conclusion_title

    for paragraph in segments.body_paragraphs:
        if paragraph._p is None:
            continue
        text = paragraph.text.strip()
        if not text:
            continue

        # 引言：前换页（自然）+ 前不空 + 后空一行；不参与编号
        if paragraph is intro_title:
            _set_paragraph_text(paragraph, INTRO_TITLE_TEXT)
            _apply_page_title_style(paragraph, h1_rule)
            _remove_blanks_before(paragraph)
            _ensure_blank_after(paragraph, body_rule)
            _collapse_blanks_after(paragraph)
            chapter_idx = 0
            section_idx = 0
            sub_idx = 0
            if report is not None:
                report.section_titles += 1
            continue

        # 结语：上下各空一行；不参与编号
        if paragraph is conclusion_title:
            _set_paragraph_text(paragraph, CONCLUSION_TITLE_TEXT)
            apply_paragraph_rule(paragraph, h1_rule)
            _ensure_blank_before(paragraph, body_rule)
            _collapse_blanks_before(paragraph)
            _ensure_blank_after(paragraph, body_rule)
            _collapse_blanks_after(paragraph)
            if report is not None:
                report.section_titles += 1
            continue

        # 脚注 / 图片：保持原样
        if paragraph._p.xpath(".//w:footnoteReference"):
            continue
        if paragraph._p.xpath(".//w:drawing"):
            continue

        if _is_figure_caption(text):
            _apply_figure_caption(paragraph, body_rule)
            if report is not None:
                report.figure_captions += 1
            continue

        if _looks_like_heading(paragraph):
            level = _resolve_heading_level(paragraph)
            plain = _strip_number_prefix(text)
            if level == 1:
                chapter_idx += 1
                section_idx = 0
                sub_idx = 0
                _set_paragraph_text(paragraph, plain)
                _set_heading_numbering(paragraph, heading_num_id, 0)
                apply_paragraph_rule(paragraph, h1_rule)
                _ensure_blank_before(paragraph, body_rule)
                _collapse_blanks_before(paragraph)
                _ensure_blank_after(paragraph, body_rule)
                _collapse_blanks_after(paragraph)
                if report is not None:
                    report.body_titles += 1
            elif level == 2:
                if chapter_idx == 0:
                    chapter_idx = 1
                section_idx += 1
                sub_idx = 0
                _set_paragraph_text(paragraph, plain)
                _set_heading_numbering(paragraph, heading_num_id, 1)
                apply_paragraph_rule(paragraph, h2_rule)
                # 二级标题上下无换行
                _remove_blanks_before(paragraph)
                _remove_blanks_after(paragraph)
                if report is not None:
                    report.body_titles += 1
            else:
                if chapter_idx == 0:
                    chapter_idx = 1
                if section_idx == 0:
                    section_idx = 1
                sub_idx += 1
                _set_paragraph_text(paragraph, plain)
                _set_heading_numbering(paragraph, heading_num_id, 2)
                apply_paragraph_rule(paragraph, h3_rule)
                # 三级标题：上空一行，下不空
                _ensure_blank_before(paragraph, body_rule)
                _collapse_blanks_before(paragraph)
                _remove_blanks_after(paragraph)
                if report is not None:
                    report.body_titles += 1
            continue

        # 普通正文：不清理前后空行（保留手动空行）
        apply_paragraph_rule(paragraph, body_rule)
        if report is not None:
            report.body_text += 1


# ---------- References ----------

def process_references(segments: DocumentSegments, config: dict, report: FormatReport | None = None) -> None:
    """独立处理参考文献段。"""
    if segments.references_title is None:
        return

    title = segments.references_title
    _set_paragraph_text(title, REFERENCES_TITLE_TEXT)
    _apply_page_title_style(title, config["styles"][STYLE_H1])
    if report is not None:
        report.section_titles += 1

    # 前换页、前不空行、后空一行
    _ensure_page_break_before(title)
    _remove_blanks_before(title)
    body_rule = config["styles"][STYLE_BODY]
    _ensure_blank_after(title, body_rule)
    _collapse_blanks_after(title)

    for paragraph in segments.references_body:
        if _is_blank_paragraph(paragraph):
            continue
        apply_paragraph_rule(paragraph, REFERENCES_BODY_RULE)
        if report is not None:
            report.references_text += 1


# ---------- Acknowledgment ----------

def process_acknowledgment(segments: DocumentSegments, config: dict, report: FormatReport | None = None) -> None:
    """独立处理致谢段。"""
    if segments.ack_title is None:
        return

    title = segments.ack_title
    _set_paragraph_text(title, ACK_TITLE_TEXT)
    _apply_page_title_style(title, config["styles"][STYLE_H1])
    if report is not None:
        report.section_titles += 1

    _ensure_page_break_before(title)
    _remove_blanks_before(title)
    body_rule = config["styles"][STYLE_BODY]
    _ensure_blank_after(title, body_rule)
    _collapse_blanks_after(title)

    for paragraph in segments.ack_body:
        if _is_blank_paragraph(paragraph):
            continue
        if paragraph._p.xpath(".//w:footnoteReference"):
            continue
        apply_paragraph_rule(paragraph, body_rule)
        if report is not None:
            report.acknowledgment_text += 1


# ============================================================
# --- document-wide setup ---
# ============================================================

def _ensure_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _section_index_for_paragraph(document: Document, paragraph) -> int:
    target_p = paragraph._p
    section_idx = 0
    body = document._body._element
    for child in body.iterchildren():
        if child is target_p:
            return min(section_idx, len(document.sections) - 1)
        if child.tag != qn("w:p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
        if sect_pr is not None:
            section_idx += 1
    return min(section_idx, len(document.sections) - 1)


def _ensure_section_break_at(document: Document, paragraph) -> "object":
    """保证 paragraph 是某个 section 的首段；返回该 section。"""
    prev = _prev_paragraph(paragraph)
    sections = list(document.sections)
    if prev is None:
        return sections[0]

    p_pr = prev._p.find(qn("w:pPr"))
    existing_sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
    if existing_sect_pr is not None:
        for idx, section in enumerate(sections):
            if section._sectPr is existing_sect_pr:
                # 段落级 sectPr 描述的是 *前一个* section，所以目标 section 是下一个。
                if idx + 1 < len(sections):
                    return sections[idx + 1]
                return sections[idx]
        return sections[-1]

    # 前一段没有独立 sectPr：把 body 级 sectPr 复制到前一段，形成新分节。
    body_sect_pr = document._body._element.sectPr
    if body_sect_pr is None:
        return sections[0]
    p_pr = prev._p.get_or_add_pPr()
    p_pr.append(deepcopy(body_sect_pr))
    return document.sections[_section_index_for_paragraph(document, paragraph)]


def apply_document_setup(
    document: Document,
    segments: DocumentSegments,
    config: dict,
    report: FormatReport | None = None,
) -> None:
    """全局设置：页边距、引言处分节、引言起页码。"""
    _ensure_margins(document)
    if report is not None:
        report.margins_updated = len(document.sections)
    page_number_config = config.get("page_number", {})

    # 页码起点：优先引言；没有引言则退化为首个非空段落。
    start_paragraph = segments.intro_title
    if start_paragraph is None:
        for para in document.paragraphs:
            if para.text.strip():
                start_paragraph = para
                break
    if start_paragraph is None:
        return

    intro_section = _ensure_section_break_at(document, start_paragraph)
    intro_sect_pr = intro_section._sectPr

    sections = list(document.sections)
    intro_section_idx = 0
    for idx, section in enumerate(sections):
        if section._sectPr is intro_sect_pr:
            intro_section_idx = idx
            break

    for idx, section in enumerate(sections):
        section.different_first_page_header_footer = False
        section.odd_and_even_pages_header_footer = False
        if idx < intro_section_idx:
            _remove_footer_references(section)
            continue
        footer = section.footer
        footer.is_linked_to_previous = False
        _clear_footer(footer)

    _set_page_number_start(intro_section, int(page_number_config.get("start", 1)))

    intro_reached = False
    for section in sections:
        # document.sections 每次迭代返回新的 Section 包装对象，
        # 比较底层 sectPr 节点才是稳定身份。
        if section._sectPr is intro_sect_pr:
            intro_reached = True
        if not intro_reached:
            continue
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = ""
        footer_para.alignment = (
            _resolve_alignment(page_number_config.get("alignment"))
            or WD_PARAGRAPH_ALIGNMENT.RIGHT
        )
        _add_page_number_run(footer_para)
        for run in footer_para.runs:
            _set_run_font(
                run,
                str(page_number_config.get("east_asia_font", "宋体")),
                float(page_number_config.get("font_size_pt", 9)),
                False,
            )
    if report is not None:
        report.page_numbers_updated = max(len(sections) - intro_section_idx, 0)


# ============================================================
# --- entry points ---
# ============================================================

def apply_mvp_format(doc_path: Path, output_path: Path, config: dict) -> FormatReport:
    """
    流水线入口：
      1. 扫描切段
      2. 正文（引言 -> 结语）
      3. 参考文献
      4. 致谢
      5. 全局设置（页边距 / 分节 / 页码）
    每段处理互不依赖；顺序与作者写作顺序自然一致。
    """
    document = Document(str(doc_path))
    report = FormatReport(total_paragraphs=len(document.paragraphs))
    segments = scan_segments(document, config)
    landmark_names: list[str] = []
    if segments.intro_title is not None:
        landmark_names.append("引言")
    if segments.conclusion_title is not None:
        landmark_names.append("结语")
    if segments.references_title is not None:
        landmark_names.append("参考文献")
    if segments.ack_title is not None:
        landmark_names.append("致谢")
    report.landmarks = landmark_names
    process_body(segments, config, report)
    process_references(segments, config, report)
    process_acknowledgment(segments, config, report)
    apply_document_setup(document, segments, config, report)
    document.save(str(output_path))
    return report


def iter_docx_files(root_dir: Path) -> list[Path]:
    return [
        p
        for p in root_dir.rglob("*.docx")
        if p.is_file() and not p.name.startswith("~$") and not p.stem.endswith(OUTPUT_SUFFIX)
    ]


def build_output_path(src: Path) -> Path:
    return src.with_name(f"{src.stem}{OUTPUT_SUFFIX}{src.suffix}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Format .docx files in a directory and generate copies."
    )
    parser.add_argument(
        "directory",
        nargs="?",
        type=Path,
        help="Target directory containing .docx files (optional, default from config)",
    )
    parser.add_argument(
        "--config",
        type=Path,
        default=DEFAULT_CONFIG_PATH,
        help=f"Path to config json (default: {DEFAULT_CONFIG_PATH.name})",
    )
    parser.add_argument(
        "--debug",
        action="store_true",
        help="Enable debug logs for landmark detection",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.debug:
        os.environ["PAPER_FORMAT_DEBUG"] = "1"
    try:
        config = load_config(args.config)
    except Exception as exc:
        print(f"[ERROR] Failed to load config: {exc}")
        return 1

    config_base_dir = args.config.resolve().parent
    configured_directory = Path(config.get("directory", "."))
    if not configured_directory.is_absolute():
        configured_directory = (config_base_dir / configured_directory).resolve()
    target_dir = args.directory or configured_directory
    target_dir = target_dir.resolve()

    if not target_dir.exists() or not target_dir.is_dir():
        print(f"[ERROR] Directory not found: {target_dir}")
        return 1

    docx_files = iter_docx_files(target_dir)
    if not docx_files:
        print(f"[INFO] No .docx files found in: {target_dir}")
        return 0

    for src in docx_files:
        out = build_output_path(src)
        apply_mvp_format(src, out, config)
        print(f"[OK] {src} -> {out}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
